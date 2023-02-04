from datetime import datetime, timezone
import pandas as pd
from openpyxl.worksheet.worksheet import Worksheet
from pandas import json_normalize
import requests
import feedparser
from bs4 import BeautifulSoup
import re
import calendar
import time
import configparser
from openpyxl import Workbook
from docx import Document


# Code form
# https://medium.com/analytics-vidhya/web-scraping-news-data-rss-feeds-python-and-google-cloud-platform-7a0df2bafe44

class news_feed_parser:
    # Parse rss feed url
    def parse_rss_feed(self, url):
        # Read feed xml data
        # Try 3 times requesting the url if error
        for i in range(0, 4):
            try:
                news_feed = feedparser.parse(url)
                break
            except:
                print('ERROR calling URL:', url, 'iter: ', (i + 1))
                pass

        # Flatten data
        df_news_feed = json_normalize(news_feed.entries)

        return df_news_feed

    # Process articles dataset
    def process_article(self, article_url):
        # Loop to parse each rss feed article url
        articles_dataset = pd.DataFrame(columns=['url', 'title', 'text'])
        article_text, article_title = self.parse_article(article_url)
        if article_text != None:
            new_data = pd.DataFrame([{'url': article_url
                                         , 'title': article_title
                                         , 'text': article_text
                                      }], columns=['url', 'title', 'text'])
            articles_dataset = pd.concat((articles_dataset, new_data), ignore_index=True, sort=False)

        articles_dataset['text_length'] = articles_dataset['text'].apply(lambda x: len(x))
        return articles_dataset

    # Build and return parent hierarchy
    def get_parent_hierarchy(self, article_parents):
        # Loop for each paragraph parent to extract its element name and id
        parents_list = []
        for parent in article_parents:
            # Extract the parent id attribute if it exists
            Parent_id = ''
            try:
                Parent_id = parent['id']
            except:
                pass

            # Append the parent name and id to the parents table
            parents_list.append(parent.name + 'id: ' + Parent_id)

        # 2.2 Construct paragraph parent hierarchy
        parent_element_list = ['' if (x == 'None' or x is None) else x for x in parents_list]
        parent_element_list.reverse()
        parent_hierarchy = ' -> '.join(parent_element_list)

        return parent_hierarchy

    # Parse url and return article text
    def parse_article(self, article_url) -> (str, str):
        # Request the article url to get the web page content.
        article_result = requests.get(article_url)
        article_content = BeautifulSoup(article_result.content, 'html.parser')

        # Bonus step : fetch the page title
        article_title = 'no title'
        title_tag = article_content.findAll('title')
        if len(title_tag) > 0:
            article_title = title_tag[0].text

        # 1. extract all paragraph elements inside the page body
        articles_body = article_content.findAll('body')

        p_blocks = articles_body[0].findAll('p')

        # 2. for each paragraph, construct its patents elements hierarchy
        # Create a dataframe to collect p_blocks data
        p_blocks_df = pd.DataFrame(columns=['element_name', 'parent_hierarchy', 'element_text', 'element_text_Count'])

        # 2.1 loop for each paragraph block
        article_text = ''
        for i in range(0, len(p_blocks)):
            # Get paragraph parent hierarchy
            parent_hierarchy = self.get_parent_hierarchy(p_blocks[i].parents)

            # Append p_blocks_df with the current paragraph data
            new_block_data = pd.DataFrame([{"element_name": p_blocks[i].name
                                               , "parent_hierarchy": parent_hierarchy
                                               , "element_text": p_blocks[i].text
                                               , "element_text_Count": len(str(p_blocks[i].text))
                                            }])

            p_blocks_df = pd.concat([p_blocks_df, new_block_data]
                                    , ignore_index=True
                                    , sort=False
                                    )

            # 3. concatenate paragraphs under the same parent hierarchy
            if len(p_blocks_df) > 0:
                # Group paragraphs by parent_hierarchy
                p_blocks_df_groupby_parent_hierarchy = p_blocks_df.groupby(by=['parent_hierarchy'])

                # Sum the paragraph lenght for each paragraph group
                p_blocks_df_groupby_parent_hierarchy_sum = p_blocks_df_groupby_parent_hierarchy[
                    ['element_text_Count']].sum()
                p_blocks_df_groupby_parent_hierarchy_sum.reset_index(inplace=True)

            # 4. select the longest paragraph as the main article
            max_id = p_blocks_df_groupby_parent_hierarchy_sum.loc[
                p_blocks_df_groupby_parent_hierarchy_sum['element_text_Count'].idxmax()
                , 'parent_hierarchy']
            article_text = '\n'.join(
                p_blocks_df.loc[p_blocks_df['parent_hierarchy'] == max_id, 'element_text'].to_list())

        # Return article text
        return article_text, article_title


def parse_mail_extract(extract_file: str) -> list:
    extract_folder = "data/0_mailExtract/"
    f = open(extract_folder + extract_file, "r")
    raw_text = f.read()

    # Remove the forced newline that truncates lines
    raw_text = raw_text.replace("=\r\n", "")
    raw_text = raw_text.replace("=\n", "")
    # Replace unicode '&' with '&'
    raw_text = raw_text.replace("\\u0026", "&")
    # Replace encoded '&' with '&'
    raw_text = raw_text.replace("&amp;", "&")

    p = re.compile('&url=3D(.*?)&ct')
    urls = p.findall(raw_text)

    # Remove duplicates
    unique_urls = list(dict.fromkeys(urls))

    # Build file name
    current_GMT = time.gmtime()
    time_stamp = calendar.timegm(current_GMT)
    url_list_folder = "data/1_urlList/"
    result_filename = "urls_list" + str(time_stamp) + ".txt"

    f = open(url_list_folder + result_filename, "w")
    f.writelines(url + '\n' for url in unique_urls)

    f.close()

    return unique_urls


class EmptyTextException(Exception):
    """Raised when the article text is empty"""
    pass


def init_sheet():
    """
    Init the Workbook, a worksheet for the result and a worksheet for the error
    :return: [Workbook, Worksheet, Worksheet]
    """

    # Create Workbook
    workbook = Workbook()
    ws_result: Worksheet = workbook.active
    ws_result.title = "Results"
    ws_error: Worksheet = workbook.create_sheet("Errors")

    # Add Title to column in the first row for the Results sheet
    ws_result['A1'] = 'URL'
    ws_result['B1'] = 'Title'
    ws_result['C1'] = 'Text'

    # Add Title to column in the first row for the Error sheet
    ws_error['A1'] = 'Original URL'
    ws_error['B1'] = 'Working URL'
    ws_error['C1'] = 'Reason'

    return workbook, ws_result, ws_error


def add_row(work_sheet, row_index, dataset):
    working_url = dataset.url[0]
    title = dataset.title[0]
    text = dataset.text[0]

    if not text:
        raise EmptyTextException

    work_sheet['A' + str(row_index)] = working_url
    work_sheet['B' + str(row_index)] = title
    work_sheet['C' + str(row_index)] = text


def add_error_row(error_work_sheet, row_index, original_url, working_url, reason='Generic Error'):
    error_work_sheet['A' + str(row_index)] = original_url
    error_work_sheet['B' + str(row_index)] = working_url
    error_work_sheet['C' + str(row_index)] = reason


def get_config():
    config = configparser.ConfigParser()
    config.read('config.ini')

    return config


def get_rss_feed_url():
    return get_config()['GoogleAlert']['rss_feed_url']


def get_max_number_of_articles_to_get():
    return int(get_config()['GoogleAlert']['max_number_of_articles_to_get'])


def build_xlsx_file():
    wb, ws, ws_error = init_sheet()

    # We use our own index instead of the for idx because if we have errors we will have
    # blank row in the Result sheet
    error_idx = 0
    result_idx = 0
    idx = 0
    for url in urls:
        idx = idx + 1
        print("Working on news #" + str(idx) + " of " + str(number_of_urls))

        try:
            dataset = my_rssFeed.process_article(url)
            add_row(ws, result_idx + 2, dataset)
            result_idx += 1
        except IndexError:
            add_error_row(ws_error, error_idx + 2, url, url)
            error_idx += 1
            print("Error while getting: {} ".format(url))
            continue
        except EmptyTextException:
            reason_text = 'Empty article text'
            add_error_row(ws_error, error_idx + 2, url, url, reason_text)
            error_idx += 1
            print("Error: {} while getting: {} ".format(reason_text, url))
            continue
        except:
            print("Error: File Upload failed.")
            continue

    # Build file name
    current_GMT = time.gmtime()
    time_stamp = calendar.timegm(current_GMT)
    filename = "data/2_results/" + "results_" + str(time_stamp) + ".xlsx"
    wb.save(filename)


def add_paragraph(document, dataset):
    working_url = dataset.url[0]
    title = dataset.title[0]
    text = dataset.text[0]

    if not text:
        raise EmptyTextException

    document.add_heading(title, level=1)
    document.add_paragraph(working_url)
    document.add_paragraph(text)


def add_error_paragraph(document, url, idx, reason='Generic Error'):
    document.add_paragraph(str(idx) + ' ' + url + ' -> Raison: ' + reason)


def build_docx_file():
    current_GMT = time.gmtime()
    time_stamp = calendar.timegm(current_GMT)

    # Document to store results
    document: Document = Document()
    document.add_heading('Résultats de l\'export du ' + str(time_stamp), 0)

    # Document to store error
    error_document: Document = Document()
    error_document.add_heading('Erreurs de l\'export du ' + str(time_stamp), 0)

    # We use our own index instead of the for idx because if we have errors we will have
    # blank row in the Result sheet
    idx = 0
    for url in urls:
        idx = idx + 1
        print("Working on news #" + str(idx) + " of " + str(number_of_urls))

        try:
            dataset = my_rssFeed.process_article(url)
            add_paragraph(document, dataset)
        except IndexError as e:
            reason_text = 'Index Error'
            add_error_paragraph(error_document, url, idx, reason_text)
            print("Error while getting: {} ".format(url))
            print(e)
            continue
        except EmptyTextException:
            reason_text = 'Empty article text'
            add_error_paragraph(error_document, url, idx, reason_text)
            print("Error: {} while getting: {} ".format(reason_text, url))
            continue
        except Exception as e:
            add_error_paragraph(error_document, url, idx)
            print("Generic Error")
            print(e)
            continue

    # Build file name
    filename = "data/2_results/" + "results_" + str(time_stamp) + ".docx"
    document.save(filename)

    error_filename = "data/2_results/" + "results_" + str(time_stamp) + "_error" + ".docx"
    error_document.save(error_filename)


if __name__ == '__main__':
    urls = parse_mail_extract("extract.mbox")

    number_of_urls = len(urls)
    print("Info: Found " + str(number_of_urls) + " to get.")

    # Applying limit to the number of article to fetch if needed
    max_article_limit = get_max_number_of_articles_to_get()
    if number_of_urls > max_article_limit:
        print("Info: Limit of article to get set to " + str(max_article_limit) +
              " in the Config file. Articles after the limit will be ignored.")
        urls = urls[0:max_article_limit]
        number_of_urls = len(urls)

    my_rssFeed = news_feed_parser()

    build_docx_file()

    # url = get_rss_feed_url()
    #
    # # Try running the search-and-extract process
    # try:
    #     # Get articles from news feed
    #     my_rssFeed = news_feed_parser()
    #     news_feed_data = my_rssFeed.parse_rss_feed(url)
    # except:
    #     print("Error: calling news feed url: {} ".format(url))
    #     raise

    # # Get number of items to work on
    # number_of_news = len(news_feed_data)
    # print('news_feed_data len: ', number_of_news)
    #
    # wb, ws, ws_error = init_sheet()
    #
    # # We use our own index instead of the for idx because if we have errors we will have
    # # blank row in the Result sheet
    # error_idx = 0
    # result_idx = 0
    # for idx, row in news_feed_data.iterrows():
    #     print("Working on news #" + str(idx) + " of " + str(number_of_news))
    #
    #     try:
    #         link = row.link
    #         re_result = re.search("^.*url=(.*)&ct.*$", link)
    #         true_link = re_result.group(1)
    #
    #         try:
    #             dataset = my_rssFeed.process_article(true_link, row.title)
    #             add_row(ws, result_idx + 2, dataset)
    #             result_idx += 1
    #         except IndexError:
    #             add_error_row(ws_error, error_idx + 2, link, true_link)
    #             error_idx += 1
    #             print("Error while getting: {} ".format(link))
    #             continue
    #         except EmptyTextException:
    #             reason_text = 'Empty article text'
    #             add_error_row(ws_error, error_idx + 2, link, true_link, reason_text)
    #             error_idx += 1
    #             print("Error: {} while getting: {} ".format(reason_text, link))
    #             continue
    #
    #     except:
    #         print("Error: File Upload failed.")
    #         continue
